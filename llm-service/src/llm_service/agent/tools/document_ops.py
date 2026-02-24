"""Document generation tools - thesis template creation."""

import logging
from pathlib import Path
from typing import Optional

from .types import ToolResult

logger = logging.getLogger(__name__)


class DocumentOpsMixin:
    """Mixin providing document generation tools."""

    def create_thesis_template(
        self,
        output_path: str = "thesis_template.docx",
    ) -> ToolResult:
        """
        Create a Chinese master's degree thesis DOCX template.

        Generates a fully formatted template following the standard
        Chinese university format (GB/T 7714-2015), including:
        cover page, declaration, abstracts, TOC, 5 chapters,
        references, acknowledgements, and author bio.

        Args:
            output_path: Where to save the .docx file

        Returns:
            ToolResult with the path to the generated file
        """
        try:
            from docx import Document  # noqa: F811
        except ImportError:
            return ToolResult(
                False, "",
                "python-docx is not installed. Run: pip install python-docx"
            )

        resolved = self._resolve_path(output_path)
        resolved.parent.mkdir(parents=True, exist_ok=True)

        try:
            # Import the standalone generator
            import importlib.util
            script = Path(__file__).resolve().parents[3] / "../../scripts/create_thesis_template.py"

            # If the script isn't available, use inline generation
            if not script.exists():
                return self._generate_thesis_inline(resolved)

            spec = importlib.util.spec_from_file_location(
                "create_thesis_template", str(script)
            )
            mod = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(mod)
            mod.create_thesis_template(str(resolved))

            return ToolResult(
                True,
                f"Thesis template created: {resolved}",
                data={"path": str(resolved)},
            )
        except Exception:
            # Fallback to inline generation
            return self._generate_thesis_inline(resolved)

    def _generate_thesis_inline(self, resolved: Path) -> ToolResult:
        """Inline fallback thesis generator (no external script needed)."""
        try:
            from docx import Document
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            from docx.oxml.ns import qn

            doc = Document()

            # Page setup
            sec = doc.sections[0]
            sec.page_width = Cm(21.0)
            sec.page_height = Cm(29.7)
            sec.top_margin = Cm(2.54)
            sec.bottom_margin = Cm(2.54)
            sec.left_margin = Cm(3.0)
            sec.right_margin = Cm(2.54)

            def _add(text, cn="宋体", sz=12, bold=False,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY):
                p = doc.add_paragraph()
                p.alignment = align
                r = p.add_run(text)
                r.font.name = "Times New Roman"
                r.font.size = Pt(sz)
                r.font.bold = bold
                r._element.rPr.rFonts.set(qn('w:eastAsia'), cn)
                return p

            # Cover
            _add("X X 大 学", sz=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            _add("硕 士 学 位 论 文", cn="黑体", sz=22, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
            for label in ["论文题目：（在此填写）", "学科专业：（专业名称）",
                          "研 究 生：（姓名）", "指导教师：（导师姓名）"]:
                _add(label, sz=14, align=WD_ALIGN_PARAGRAPH.CENTER)
            doc.add_page_break()

            # Declaration
            _add("独创性声明", cn="黑体", sz=16, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
            _add("本人声明所呈交的学位论文是本人在导师指导下进行的研究工作"
                 "及取得的研究成果……", sz=12)
            doc.add_page_break()

            # Chinese abstract
            _add("摘  要", cn="黑体", sz=16, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
            _add("（在此撰写中文摘要，约300-500字……）", sz=12)
            doc.add_page_break()

            # English abstract
            _add("ABSTRACT", sz=16, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
            _add("(Write English abstract here, 200-300 words...)", sz=12)
            doc.add_page_break()

            # TOC placeholder
            _add("目  录", cn="黑体", sz=16, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
            _add("（请在Word中生成目录）", sz=12)
            doc.add_page_break()

            # Chapters
            chapters = [
                ("第一章  绪论", ["1.1 研究背景", "1.2 研究意义",
                               "1.3 国内外研究现状", "1.4 研究内容"]),
                ("第二章  相关理论与技术", ["2.1 基本概念", "2.2 相关理论",
                                       "2.3 关键技术"]),
                ("第三章  研究方法与设计", ["3.1 总体方案", "3.2 模型设计",
                                       "3.3 关键问题"]),
                ("第四章  实现与实验", ["4.1 实验环境", "4.2 实现过程",
                                    "4.3 实验结果"]),
                ("第五章  总结与展望", ["5.1 工作总结", "5.2 不足与展望"]),
            ]
            for ch_title, sections in chapters:
                _add(ch_title, cn="黑体", sz=16, bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER)
                for sec_title in sections:
                    _add(sec_title, cn="黑体", sz=14, bold=True,
                         align=WD_ALIGN_PARAGRAPH.LEFT)
                    _add("（在此撰写内容……）", sz=12)

            # References
            doc.add_page_break()
            _add("参考文献", cn="黑体", sz=16, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
            _add("[1]  作者. 标题[J]. 期刊, 年, 卷(期): 页码.", sz=10.5)

            # Acknowledgements
            doc.add_page_break()
            _add("致  谢", cn="黑体", sz=16, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
            _add("（在此撰写致谢内容……）", sz=12)

            doc.save(str(resolved))
            return ToolResult(
                True,
                f"Thesis template created: {resolved}",
                data={"path": str(resolved)},
            )
        except Exception as e:
            return ToolResult(False, "", f"Failed to generate thesis: {e}")
