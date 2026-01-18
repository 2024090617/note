"""Document generation with python-docx and LLM-based formatting."""

from typing import List, Optional, Dict, Any
import json
from pathlib import Path
import logging

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
except ImportError:
    Document = None  # Will be installed later

from ..client import LLMClient
from .models import FormattingRules, ThesisSection, CitationStyle
from .prompts import FORMATTING_INTERPRETATION_PROMPT

logger = logging.getLogger(__name__)


class FormattingInterpreter:
    """Use LLM to interpret natural language formatting specifications."""
    
    def __init__(self, llm_client: Optional[LLMClient] = None):
        """Initialize formatting interpreter.
        
        Args:
            llm_client: LLM client for interpretation
        """
        self.llm = llm_client or LLMClient()
    
    def interpret(self, user_spec: str) -> FormattingRules:
        """Interpret natural language formatting spec.
        
        Args:
            user_spec: User's formatting description (e.g., "标准中国高校硕士论文")
            
        Returns:
            Structured formatting rules
        """
        prompt = FORMATTING_INTERPRETATION_PROMPT.format(user_spec=user_spec)
        
        try:
            response = self.llm.simple_query(
                prompt=prompt,
                temperature=0.3  # Lower temperature for consistency
            )

            logger.debug(f"Formatting interpretation response: {response}")
            
            # Extract JSON from response
            json_str = self._extract_json(response)
            rules_dict = json.loads(json_str)
            
            # Add original spec
            rules_dict["user_specification"] = user_spec
            
            # Parse into FormattingRules
            rules = FormattingRules(**rules_dict)
            
            return rules
            
        except Exception as e:
            logger.warning(f"Failed to interpret formatting spec, using defaults: {e}")
            return FormattingRules(user_specification=user_spec)
    
    def _extract_json(self, text: str) -> str:
        """Extract JSON from LLM response.
        
        Args:
            text: Response text
            
        Returns:
            JSON string
        """
        # Try to find JSON object
        start = text.find('{')
        end = text.rfind('}')
        
        if start != -1 and end != -1:
            return text[start:end+1]
        
        return text


class ThesisDocxGenerator:
    """Generate .docx documents with proper formatting."""
    
    def __init__(self, formatting: FormattingRules):
        """Initialize document generator.
        
        Args:
            formatting: Formatting rules to apply
        """
        if Document is None:
            raise ImportError(
                "python-docx is not installed. "
                "Install it with: pip install python-docx"
            )
        
        self.formatting = formatting
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """Set up document-level formatting."""
        # Set page margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = self._parse_size(self.formatting.page_margins.get("top", "2.54cm"))
            section.bottom_margin = self._parse_size(self.formatting.page_margins.get("bottom", "2.54cm"))
            section.left_margin = self._parse_size(self.formatting.page_margins.get("left", "3.17cm"))
            section.right_margin = self._parse_size(self.formatting.page_margins.get("right", "3.17cm"))
        
        # Set up styles
        self._setup_styles()
    
    def _setup_styles(self):
        """Set up paragraph and character styles."""
        styles = self.doc.styles
        
        # Body style
        if "Body" in [s.name for s in styles]:
            body_style = styles['Body']
        else:
            body_style = styles.add_style('Body', WD_STYLE_TYPE.PARAGRAPH)
        
        # Font
        font = body_style.font
        font.name = self.formatting.body_font.get("english", "Times New Roman")
        font.size = self._parse_font_size(self.formatting.body_font_size)
        
        # Chinese font
        body_style.element.rPr.rFonts.set(qn('w:eastAsia'), self.formatting.body_font.get("chinese", "宋体"))
        
        # Paragraph formatting
        pf = body_style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = self.formatting.line_spacing
        pf.first_line_indent = self._parse_size(self.formatting.first_line_indent)
        pf.alignment = self._parse_alignment(self.formatting.body_alignment)
        
        # Heading styles - 按照中国高校硕士论文标准
        # 一级标题：宋体三号(16pt)，加粗，居中
        heading1_name = 'Heading 1'
        if heading1_name in [s.name for s in styles]:
            h1_style = styles[heading1_name]
        else:
            h1_style = styles.add_style(heading1_name, WD_STYLE_TYPE.PARAGRAPH)
        
        h1_font = h1_style.font
        h1_font.name = "宋体"
        h1_font.size = Pt(16)  # 三号
        h1_font.bold = True
        h1_style.element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        
        h1_pf = h1_style.paragraph_format
        h1_pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1_pf.space_before = Pt(12)
        h1_pf.space_after = Pt(12)
        
        # 二级标题：宋体四号(14pt)，加粗，左对齐
        heading2_name = 'Heading 2'
        if heading2_name in [s.name for s in styles]:
            h2_style = styles[heading2_name]
        else:
            h2_style = styles.add_style(heading2_name, WD_STYLE_TYPE.PARAGRAPH)
        
        h2_font = h2_style.font
        h2_font.name = "宋体"
        h2_font.size = Pt(14)  # 四号
        h2_font.bold = True
        h2_style.element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        
        h2_pf = h2_style.paragraph_format
        h2_pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2_pf.space_before = Pt(6)
        h2_pf.space_after = Pt(6)
        
        # 三级标题：宋体小四号(12pt)，加粗，左对齐
        heading3_name = 'Heading 3'
        if heading3_name in [s.name for s in styles]:
            h3_style = styles[heading3_name]
        else:
            h3_style = styles.add_style(heading3_name, WD_STYLE_TYPE.PARAGRAPH)
        
        h3_font = h3_style.font
        h3_font.name = "宋体"
        h3_font.size = Pt(12)  # 小四号
        h3_font.bold = True
        h3_style.element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        
        h3_pf = h3_style.paragraph_format
        h3_pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h3_pf.space_before = Pt(6)
        h3_pf.space_after = Pt(6)
    
    def add_cover_page(self, title: str, author: str, **kwargs):
        """Add cover page.
        
        Args:
            title: Thesis title
            author: Author name
            **kwargs: Additional info (institution, date, advisor, etc.)
        """
        # Add some spacing at top
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Title - 黑体二号(22pt)，加粗，居中
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(22)  # 二号
        run.font.bold = True
        run.font.name = "Times New Roman"
        run.element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
        
        # Spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Author - 宋体三号(16pt)
        if author:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"作者：{author}")
            run.font.size = Pt(16)  # 三号
            run.font.name = "Times New Roman"
            run.element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        
        # Institution - 宋体三号(16pt)
        if "institution" in kwargs:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(kwargs["institution"])
            run.font.size = Pt(16)  # 三号
            run.font.name = "Times New Roman"
            run.element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        
        # Advisor - 宋体三号(16pt)
        if "advisor" in kwargs:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"指导教师：{kwargs['advisor']}")
            run.font.size = Pt(16)  # 三号
            run.font.name = "Times New Roman"
            run.element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        
        self.doc.add_paragraph()
        
        # Date - 宋体四号(14pt)
        if "date" in kwargs:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(kwargs["date"])
            run.font.size = Pt(14)  # 四号
            run.font.name = "Times New Roman"
            run.element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        
        # Page break
        self.doc.add_page_break()
    
    def add_toc(self):
        """Add table of contents."""
        if not self.formatting.include_toc:
            return
        
        # TOC title
        p = self.doc.add_paragraph(self.formatting.toc_title, style='Heading 1')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add TOC field (Note: python-docx doesn't support automatic TOC generation,
        # user needs to update fields in Word after opening)
        p = self.doc.add_paragraph()
        run = p.add_run()
        fldChar1 = run._element
        
        # Add instruction text
        instrText = self.doc.add_paragraph()
        instrText.add_run('（请在Word中右键点击此处选择"更新域"以生成目录）').italic = True
        
        self.doc.add_page_break()
    
    def add_section(self, section: ThesisSection, level: int = 1):
        """Add a thesis section.
        
        Args:
            section: Section data
            level: Heading level (1, 2, or 3)
        """
        # Section heading
        heading_text = f"{section.section_id} {section.title}"
        self.doc.add_heading(heading_text, level=level)
        
        # Section content
        paragraphs = section.content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                p = self.doc.add_paragraph(para_text, style='Body')
    
    def add_text(self, text: str, style: str = 'Body'):
        """Add text with specified style.
        
        Args:
            text: Text to add
            style: Style name
        """
        self.doc.add_paragraph(text, style=style)
    
    def add_references(self, references: List[str]):
        """Add references section.
        
        Args:
            references: List of formatted reference strings
        """
        # References heading
        self.doc.add_page_break()
        self.doc.add_heading(self.formatting.references_title, level=1)
        
        # Add each reference
        for ref in references:
            p = self.doc.add_paragraph(ref, style='Body')
            p.paragraph_format.first_line_indent = Pt(0)  # No indent for references
            p.paragraph_format.left_indent = Pt(0)
    
    def save(self, filepath: str):
        """Save document to file.
        
        Args:
            filepath: Output file path
        """
        self.doc.save(filepath)
        logger.info(f"Document saved to {filepath}")
    
    @staticmethod
    def _parse_size(size_str: str):
        """Parse size string to python-docx unit.
        
        Args:
            size_str: Size string (e.g., "2.54cm", "1in", "2em")
            
        Returns:
            Size in EMU (English Metric Units)
        """
        size_str = size_str.lower().strip()
        
        if "cm" in size_str:
            value = float(size_str.replace("cm", ""))
            return Cm(value)
        elif "in" in size_str:
            value = float(size_str.replace("in", ""))
            return Inches(value)
        elif "pt" in size_str:
            value = float(size_str.replace("pt", ""))
            return Pt(value)
        elif "em" in size_str:
            # Approximate: 1em ≈ 12pt for Chinese text
            value = float(size_str.replace("em", ""))
            return Pt(value * 12)
        else:
            # Try to parse as number (assume pt)
            try:
                value = float(size_str)
                return Pt(value)
            except ValueError:
                return Pt(12)  # Default
    
    @staticmethod
    def _parse_font_size(size_str: str):
        """Parse font size to Points.
        
        Args:
            size_str: Font size (e.g., "12pt", "小四")
            
        Returns:
            Font size in Points
        """
        size_str = size_str.lower().strip()
        
        # Chinese font sizes
        chinese_sizes = {
            "初号": 42,
            "小初": 36,
            "一号": 26,
            "小一": 24,
            "二号": 22,
            "小二": 18,
            "三号": 16,
            "小三": 15,
            "四号": 14,
            "小四": 12,
            "五号": 10.5,
            "小五": 9,
        }
        
        if size_str in chinese_sizes:
            return Pt(chinese_sizes[size_str])
        elif "pt" in size_str:
            value = float(size_str.replace("pt", ""))
            return Pt(value)
        else:
            try:
                value = float(size_str)
                return Pt(value)
            except ValueError:
                return Pt(12)  # Default
    
    @staticmethod
    def _parse_alignment(align_str: str):
        """Parse alignment string.
        
        Args:
            align_str: Alignment ("left", "center", "right", "justified")
            
        Returns:
            WD_ALIGN_PARAGRAPH enum
        """
        align_str = align_str.lower().strip()
        
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        
        return alignment_map.get(align_str, WD_ALIGN_PARAGRAPH.JUSTIFY)


def create_thesis_document(
    sections: List[ThesisSection],
    references: List[str],
    formatting_spec: str,
    title: str,
    author: str,
    output_path: str,
    **kwargs
) -> str:
    """High-level function to create complete thesis document.
    
    Args:
        sections: List of thesis sections
        references: List of formatted references
        formatting_spec: Natural language formatting specification
        title: Thesis title
        author: Author name
        output_path: Output file path
        **kwargs: Additional metadata (institution, advisor, date, etc.)
        
    Returns:
        Path to generated document
    """
    # Interpret formatting
    interpreter = FormattingInterpreter()
    formatting = interpreter.interpret(formatting_spec)
    
    # Create document
    generator = ThesisDocxGenerator(formatting)
    
    # Add cover page
    generator.add_cover_page(title, author, **kwargs)
    
    # Add TOC
    generator.add_toc()
    
    # Add sections
    for section in sections:
        # Determine heading level from section_id
        level = section.section_id.count('.') + 1
        level = min(level, 3)  # Max level 3
        generator.add_section(section, level=level)
    
    # Add references
    if references:
        generator.add_references(references)
    
    # Save
    generator.save(output_path)
    
    return output_path
